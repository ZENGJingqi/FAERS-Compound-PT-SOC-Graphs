from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "expert_package_cleaning_and_core_graphs_2026-04-22"
ASSET_DIR = PACKAGE / "_doc_assets"
OUT_DOC = PACKAGE / "FAERS_预处理与核心异构图数据说明书_2026-04-23.docx"

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Arial Unicode MS", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_chinese_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: int | None = None, bold: bool | None = None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None, size: int = 11, bold: bool = False, east_asia: str = "宋体"):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_chinese_font(run, east_asia=east_asia, size=size, bold=bold)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_chinese_font(run, size=11)


def add_table(doc: Document, df: pd.DataFrame, col_width_inches: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(col))
        set_chinese_font(run, east_asia="黑体", size=10, bold=True)
        set_cell_shading(hdr[i], "D9E2F3")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_chinese_font(run, size=10)
    if col_width_inches:
        for row in table.rows:
            for i, width in enumerate(col_width_inches):
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_chinese_font(run, east_asia="黑体", size=10)


def plot_workflow(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 2.8))
    ax.axis("off")
    labels = [
        "Step1\n病例去重",
        "Step2\n病例-药物/反应基础表",
        "Step3\n药物名称标准化\nRxNorm/DrugBank",
        "Step4\n结构标准化\nInChIKey/SMILES",
        "Step5\n反应标准化\nMedDRA PT/SOC",
        "Step14\n裁剪核心图\nge10/ge20/ge30",
    ]
    colors = ["#B23A48", "#464F5F", "#6F4C9B", "#2C5AA0", "#0D7587", "#6B8E23"]
    xs = [0.05, 0.21, 0.38, 0.55, 0.72, 0.88]
    for x, label, color in zip(xs, labels, colors):
        rect = plt.Rectangle((x - 0.06, 0.35), 0.12, 0.3, fc=color, ec="black", lw=1.2, transform=ax.transAxes)
        ax.add_patch(rect)
        ax.text(x, 0.5, label, ha="center", va="center", color="white", fontsize=10, fontweight="bold", transform=ax.transAxes)
    for i in range(len(xs) - 1):
        ax.annotate("", xy=(xs[i + 1] - 0.07, 0.5), xytext=(xs[i] + 0.07, 0.5), xycoords=ax.transAxes,
                    arrowprops=dict(arrowstyle="->", lw=1.8, color="black"))
    fig.tight_layout()
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def plot_stage_summary(path: Path, key_numbers: dict) -> None:
    labels = ["药物词条", "映射到InChI词条", "结构化成分", "PS/SS成分", "MedDRA PT", "PS/SS成分-PT对"]
    values = [
        key_numbers["step3"]["drug_terms_total"],
        key_numbers["step3"]["terms_mapped_to_inchi"],
        key_numbers["step4"]["unique_inchikey"],
        key_numbers["step4"]["inchikey_in_psss"],
        key_numbers["step5"]["distinct_meddra_pt"],
        key_numbers["step5"]["meddra_pairs_psss"],
    ]
    fig, ax = plt.subplots(figsize=(9, 4.8))
    bars = ax.bar(range(len(labels)), values, color=["#B23A48", "#464F5F", "#6F4C9B", "#2C5AA0", "#0D7587", "#6B8E23"])
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=20, ha="right", fontsize=10, color="black")
    ax.set_ylabel("数量", fontsize=12, color="black")
    ax.tick_params(axis="y", labelsize=10, colors="black")
    ax.set_yscale("log")
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    for bar, v in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, v * 1.1, f"{v:,}", ha="center", va="bottom", fontsize=9, color="black")
    for s in ax.spines.values():
        s.set_color("black")
    fig.tight_layout()
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def plot_graph_comparison(path: Path, comp_df: pd.DataFrame, ref_summary: dict) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.4))
    plot_df = pd.concat(
        [
            pd.DataFrame(
                [
                    {
                        "label": "总数据",
                        "compound_nodes": ref_summary["compound_rows"],
                        "pt_nodes": ref_summary["reaction_rows"],
                        "soc_nodes": ref_summary["soc_rows"],
                        "ps_edges": ref_summary["ps_association_rows"],
                        "ss_edges": ref_summary["ss_association_rows"],
                    }
                ]
            ),
            comp_df.assign(label=comp_df["threshold"].map(lambda x: f">={x}"))[
                ["label", "compound_nodes", "pt_nodes", "soc_nodes", "ps_edges", "ss_edges"]
            ],
        ],
        ignore_index=True,
    )
    x = range(len(plot_df))
    labels = plot_df["label"].tolist()
    axes[0].plot(x, plot_df["compound_nodes"], marker="o", color="#B23A48", lw=2, label="compound")
    axes[0].plot(x, plot_df["pt_nodes"], marker="o", color="#2C5AA0", lw=2, label="pt")
    axes[0].plot(x, plot_df["soc_nodes"], marker="o", color="#0D7587", lw=2, label="soc")
    axes[0].set_xticks(list(x))
    axes[0].set_xticklabels(labels, fontsize=10)
    axes[0].set_ylabel("节点数", fontsize=12, color="black")
    axes[0].tick_params(axis="both", colors="black", labelsize=10)
    axes[0].legend(frameon=False, fontsize=10)
    axes[0].grid(axis="y", linestyle="--", alpha=0.3)

    width = 0.24
    axes[1].bar([i - width / 2 for i in x], plot_df["ps_edges"], width=width, color="#464F5F", label="PS边")
    axes[1].bar([i + width / 2 for i in x], plot_df["ss_edges"], width=width, color="#6F4C9B", label="SS边")
    axes[1].set_xticks(list(x))
    axes[1].set_xticklabels(labels, fontsize=10)
    axes[1].set_ylabel("边数", fontsize=12, color="black")
    axes[1].tick_params(axis="both", colors="black", labelsize=10)
    axes[1].grid(axis="y", linestyle="--", alpha=0.3)
    axes[1].legend(frameon=False, fontsize=9)
    for ax in axes:
        for s in ax.spines.values():
            s.set_color("black")
    fig.tight_layout()
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def main() -> None:
    ASSET_DIR.mkdir(exist_ok=True)
    key_numbers = json.loads((PACKAGE / "00_key_numbers.json").read_text(encoding="utf-8"))
    stage_df = pd.read_csv(PACKAGE / "01_stage_changes.csv")
    overview_df = pd.read_csv(PACKAGE / "02_data_overview.csv")
    comp_df = pd.read_csv(PACKAGE / "03_pruned_graphs" / "graph_comparison.csv")
    ref_summary = json.loads((PACKAGE / "04_cleaned_reference_tables" / "reference_tables_summary.json").read_text(encoding="utf-8"))

    workflow_png = ASSET_DIR / "workflow.png"
    stage_png = ASSET_DIR / "stage_summary.png"
    graph_png = ASSET_DIR / "graph_comparison.png"
    plot_workflow(workflow_png)
    plot_stage_summary(stage_png, key_numbers)
    plot_graph_comparison(graph_png, comp_df, ref_summary)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    add_paragraph(doc, "FAERS 预处理与核心异构图数据说明书", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, east_asia="黑体")
    add_paragraph(doc, "版本：2026-04-23", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_paragraph(doc, "范围：仅包含数据收集、清洗、标准化和三个裁剪版核心图；不包含建模和预测结果。", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    doc.add_paragraph("")
    add_paragraph(doc, "一、项目目标", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "从 FAERS 原始数据中获得可复用的结构化成分-反应数据。",
        "将药物统一到 InChIKey / SMILES，将反应统一到 MedDRA PT / primary SOC。",
        "构建三个不同频次阈值下的核心异构图，供后续研究直接使用。",
    ])

    add_paragraph(doc, "二、数据收集与版本", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "FAERS 原始数据：本地收集的季度压缩包，从 2004Q1 到 2025Q4，覆盖 AERS/FAERS 全部季度数据。",
        "RxNorm：使用本地目录 `RxNorm_full_03022026`。",
        "DrugBank：使用本地目录 `drugbank_5.1.15`。",
        "MedDRA：使用本地 `MedDRA 29.0 English` 和 `MedDRA 29.0 Chinese`。",
    ])
    add_paragraph(doc, "表1 数据收集与版本信息", size=12, bold=True, east_asia="黑体")
    source_table = pd.DataFrame([
        ["FAERS / AERS", "2004Q1-2025Q4", "FDA 季度 ASCII 压缩包", "原始病例、药物、反应来源"],
        ["RxNorm", "RxNorm_full_03022026", "本地下载目录", "药物名称规范化"],
        ["DrugBank", "5.1.15", "本地下载目录", "药物结构与标准名称对接"],
        ["MedDRA English", "29.0", "本地下载目录", "反应标准化到 PT / SOC"],
        ["MedDRA Chinese", "29.0", "本地下载目录", "反应与 SOC 中文双语补充"],
    ], columns=["数据集", "版本/时间范围", "本地来源", "用途"])
    add_table(doc, source_table, [1.6, 1.5, 2.0, 2.1])

    add_paragraph(doc, "三、数据标准与统一主键", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "统一后的成分主键为 InChIKey，反应主键为 MedDRA PT，系统层级为 primary SOC。",
        "成分结构字段统一保留为 `smiles_std` 和 `inchi_std`。",
        "反应字段统一保留为 PT 编码、PT 英文名、PT 中文名，以及主 SOC 编码与主 SOC 中英文名。",
    ])

    add_paragraph(doc, "四、整体处理流程", size=14, bold=True, east_asia="黑体")
    doc.add_picture(str(workflow_png), width=Inches(6.9))
    add_caption(doc, "图1  从原始 FAERS 到三个核心图的处理流程")

    add_paragraph(doc, "五、各步骤方法与关键结果", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "Step1：病例去重，保留最新且有效病例，共 19,961,057 条。",
        "Step2：构建 case_drug 和 case_reaction 基础表，分别得到 72,212,791 条药物记录和 59,395,334 条反应记录。",
        "Step3：药物名称标准化并连接 RxNorm / DrugBank，共有 702,719 个不同药物词条，其中 183,162 个词条成功连到 InChI；对应 53,858,885 条药物记录成功连到结构。",
        "Step4：结构标准化到 InChIKey / InChI / SMILES，得到 4,523 个唯一 InChIKey，其中 3,758 个在 PS/SS 报告中实际出现。",
        "Step5：反应标准化到 MedDRA PT / primary SOC，共 59,395,333 条反应记录成功映射到 MedDRA PT，涉及 22,643 个不同 PT。",
        "Step14：基于标准化结果构建 FAERS 核心图，并裁剪出 ge10、ge20、ge30 三个版本。",
    ])
    doc.add_picture(str(stage_png), width=Inches(6.7))
    add_caption(doc, "图2  关键标准化阶段的数量变化（纵轴为对数刻度）")

    add_paragraph(doc, "表2 关键数量摘要", size=12, bold=True, east_asia="黑体")
    overview_table = overview_df.copy()
    overview_table["value"] = overview_table["value"].map(lambda x: f"{x:,}" if str(x).isdigit() else x)
    overview_table.columns = ["指标", "数值", "说明"]
    add_table(doc, overview_table, [1.9, 1.4, 3.2])

    add_paragraph(doc, "六、核心数据表说明", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        f"全部结构化成分表：{ref_summary['compound_rows']:,} 行，包含 InChIKey、标准化 SMILES、标准化 InChI、DrugBank 名称汇总以及 ge10/ge20/ge30 覆盖标记。",
        f"全部反应双语表：{ref_summary['reaction_rows']:,} 行，包含 PT 编码、PT 中英文名、主 SOC 编码、主 SOC 中英文名以及 ge10/ge20/ge30 覆盖标记。",
        f"SOC 双语表：{ref_summary['soc_rows']} 行，对应 27 个 primary SOC。",
        f"PS 全量成分-反应关联表：{ref_summary['ps_association_rows']:,} 行。",
        f"SS 全量成分-反应关联表：{ref_summary['ss_association_rows']:,} 行。",
    ])

    add_paragraph(doc, "表3 当前保留的数据表及用途", size=12, bold=True, east_asia="黑体")
    table2 = pd.DataFrame([
        ["all_compounds_basic_info.xlsx", "全部结构化成分基础信息"],
        ["all_reactions_basic_info_bilingual.xlsx", "全部反应 PT 双语基础信息"],
        ["all_soc_basic_info_bilingual.xlsx", "27 个 SOC 双语表"],
        ["all_compound_pt_ps_associations_bilingual.csv.gz", "全部 PS 成分-反应关联表"],
        ["all_compound_pt_ss_associations_bilingual.csv.gz", "全部 SS 成分-反应关联表"],
    ], columns=["文件", "说明"])
    add_table(doc, table2, [3.8, 3.0])

    add_paragraph(doc, "七、三个裁剪版核心图", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "三个图均只保留 3 类节点：compound、pt、soc。",
        "三个图均只保留 3 类边：compound_has_pt_ps、compound_has_pt_ss、pt_belongs_to_primary_soc。",
        "三个图中的 PT 和 SOC 已补充中英文双语字段。",
        "compound 节点保留 smiles_std、inchi_std 和 inchikey，可直接用于结构读取。",
    ])
    doc.add_picture(str(graph_png), width=Inches(6.9))
    add_caption(doc, "图3  ge10、ge20、ge30 三个核心图的规模比较")

    add_paragraph(doc, "表4 三个图的基本情况", size=12, bold=True, east_asia="黑体")
    graph_table = pd.DataFrame([
        {
            "图版本": "清洗完成后总数据",
            "compound节点": ref_summary["compound_rows"],
            "pt节点": ref_summary["reaction_rows"],
            "soc节点": ref_summary["soc_rows"],
            "PS边": ref_summary["ps_association_rows"],
            "SS边": ref_summary["ss_association_rows"],
            "PT-SOC边": ref_summary["reaction_rows"],
        }
    ])
    graph_table = pd.concat(
        [
            graph_table,
            comp_df.rename(
                columns={
                    "threshold": "图版本",
                    "compound_nodes": "compound节点",
                    "pt_nodes": "pt节点",
                    "soc_nodes": "soc节点",
                    "ps_edges": "PS边",
                    "ss_edges": "SS边",
                    "pt_soc_edges": "PT-SOC边",
                }
            )[["图版本", "compound节点", "pt节点", "soc节点", "PS边", "SS边", "PT-SOC边"]],
        ],
        ignore_index=True,
    )
    graph_table["图版本"] = graph_table["图版本"].map(lambda x: f"n_reports >= {x}" if str(x).isdigit() else x)
    add_table(doc, graph_table, [1.8, 1.0, 1.0, 0.8, 1.1, 1.1, 1.1])

    add_paragraph(doc, "八、推荐使用方案", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "如果只选择一张图作为默认主图，推荐优先使用 ge20。",
        "原因是 ge20 在噪声控制、覆盖范围和文件体量之间更平衡。",
        "ge10 可作为更宽覆盖版本，ge30 可作为更严格高置信版本。",
    ])

    add_paragraph(doc, "九、专家包目录说明", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "00_key_numbers.json：关键数字摘要。",
        "01_stage_changes.csv：各标准化阶段的数量变化。",
        "02_data_overview.csv：总体数据概况。",
        "03_pruned_graphs：三个裁剪版核心图及其摘要。",
        "04_cleaned_reference_tables：全部基础表和全量关联表。",
        "05_pipeline_code：最终保留的预处理与图构建脚本。",
        "06_pipeline_notes：各步骤的执行记录和方法说明。",
    ])

    add_paragraph(doc, "十、交付结论", size=14, bold=True, east_asia="黑体")
    add_bullets(doc, [
        "当前交付已经完成从 FAERS 原始数据到结构化成分-反应-SOC 图数据的完整预处理闭环。",
        "当前交付不包含任何建模和预测内容，只保留数据清洗、标准化和核心图。",
        "读者可以直接基于专家包中的 Word 说明书、基础表和 SQLite 图数据库理解并复现整个预处理过程。",
    ])

    # 设置页眉页脚安全分节
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_paragraph(doc, "附录：主要文件清单", size=14, bold=True, east_asia="黑体")
    appendix = pd.DataFrame([
        ["03_pruned_graphs/compound_pt_role_core_ge10/step14_compound_pt_role_core_ge10.sqlite", "裁剪图 ge10"],
        ["03_pruned_graphs/compound_pt_role_core_ge20/step14_compound_pt_role_core_ge20.sqlite", "裁剪图 ge20"],
        ["03_pruned_graphs/compound_pt_role_core_ge30/step14_compound_pt_role_core_ge30.sqlite", "裁剪图 ge30"],
        ["all_compounds_basic_info.xlsx", "全部结构化成分表"],
        ["all_reactions_basic_info_bilingual.xlsx", "全部反应双语表"],
        ["all_soc_basic_info_bilingual.xlsx", "全部 SOC 双语表"],
        ["04_cleaned_reference_tables/all_compound_pt_ps_associations_bilingual.csv.gz", "全部 PS 关联表"],
        ["04_cleaned_reference_tables/all_compound_pt_ss_associations_bilingual.csv.gz", "全部 SS 关联表"],
    ], columns=["文件", "说明"])
    add_table(doc, appendix, [4.8, 2.0])

    doc.save(str(OUT_DOC))
    print(OUT_DOC)


if __name__ == "__main__":
    main()
